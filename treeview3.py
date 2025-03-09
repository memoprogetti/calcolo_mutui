import streamlit as st
import pandas as pd
from docx import Document
import os

# Funzione per caricare il file Excel
def load_excel(file):
    df = pd.read_excel(file)
    return df

# Funzione per esportare in Word
def export_to_word(data, file_path):
    doc = Document()
    doc.add_heading('Dati esportati', 0)
    for item in data:
        doc.add_paragraph(str(item))
    doc.save(file_path)

# Funzione per creare la cartella di esportazione se non esiste
def create_export_folder():
    export_folder = "c:\\scambiodati_export"
    if not os.path.exists(export_folder):
        os.makedirs(export_folder)
    return export_folder

# Funzione principale dell'applicazione
def main():
    st.title("Netcom Centro Studi")

    # Usa st.file_uploader per selezionare il file
    uploaded_file = st.file_uploader("Seleziona un file Excel", type=["xlsx"])

    if uploaded_file is not None:
        df = load_excel(uploaded_file)
        
        # Mostra l'intestazione delle colonne
        st.write("Intestazione colonne:")
        st.write(df.columns.tolist())

        # Seleziona la prima colonna
        first_column = df.columns[0]
        st.write(f"Seleziona un valore dalla colonna '{first_column}':")
        selected_first = st.selectbox("", ["Tutti i record"] + df[first_column].unique().tolist())

        # Filtra i dati in base alla selezione della prima colonna
        if selected_first == "Tutti i record":
            filtered_df_1 = df
        else:
            filtered_df_1 = df[df[first_column] == selected_first]
        
        # Mostra il numero di record trovati
        st.markdown(f"<h2 style='color:red;'>Record trovati: {len(filtered_df_1)}</h2>", unsafe_allow_html=True)

        # Seleziona la seconda colonna
        second_column = df.columns[1]
        st.write(f"Seleziona un valore dalla colonna '{second_column}':")
        selected_second = st.selectbox("", ["Tutti i record"] + filtered_df_1[second_column].unique().tolist())

        # Filtra i dati in base alla selezione della seconda colonna
        if selected_second == "Tutti i record":
            filtered_df_2 = filtered_df_1
        else:
            filtered_df_2 = filtered_df_1[filtered_df_1[second_column] == selected_second]
        
        # Mostra il numero di record trovati
        st.markdown(f"<h2 style='color:red;'>Record trovati: {len(filtered_df_2)}</h2>", unsafe_allow_html=True)

        # Seleziona la terza colonna
        third_column = df.columns[2]
        st.write(f"Seleziona un valore dalla colonna '{third_column}':")
        selected_third = st.selectbox("", ["Tutti i record"] + filtered_df_2[third_column].unique().tolist())

        # Filtra i dati in base alla selezione della terza colonna
        if selected_third == "Tutti i record":
            filtered_df_3 = filtered_df_2
        else:
            filtered_df_3 = filtered_df_2[filtered_df_2[third_column] == selected_third]
        
        # Mostra il numero di record trovati
        st.markdown(f"<h2 style='color:red;'>Record trovati: {len(filtered_df_3)}</h2>", unsafe_allow_html=True)

        # Mostra i dati della colonna 4
        fourth_column = df.columns[3]
        st.write(f"Dati della colonna '{fourth_column}':")

        # Aggiungi una casella di ricerca per filtrare i dati
        search_term = st.text_input("Cerca nella colonna 4:")
        if search_term:
            filtered_df_4 = filtered_df_3[filtered_df_3[fourth_column].str.contains(search_term, case=False, na=False)]
        else:
            filtered_df_4 = filtered_df_3

        # Visualizzazione tabellare con righe alternate di colore
        st.write("Visualizzazione tabellare:")
        st.dataframe(
            filtered_df_4.style.apply(lambda x: ['background-color: #e6ffe6' if i % 2 == 0 else 'background-color: #ffffe6' for i in range(len(x))], axis=0)
        )

        # Opzioni per esportare i dati
        st.write("Esporta i dati:")
        export_folder = create_export_folder()

        if st.button("Esporta in Word"):
            file_path = os.path.join(export_folder, "exported_data.docx")
            export_to_word(filtered_df_4[fourth_column], file_path)
            st.write(f"File Word esportato con successo in: {file_path}")

        if st.button("Esporta in Excel"):
            file_path = os.path.join(export_folder, "exported_data.xlsx")
            filtered_df_4.to_excel(file_path, index=False)
            st.write(f"File Excel esportato con successo in: {file_path}")

if __name__ == "__main__":
    main()